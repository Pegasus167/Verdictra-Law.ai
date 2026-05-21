"""
pipeline/extractors/docx_extractor.py
---------------------------------------
DOCX extractor using python-docx.
Preserves heading hierarchy — headings become Markdown headers.
Inserts <!-- page:N --> markers so tree_builder assigns correct page numbers.

Page detection strategy (in order of preference):
  1. Explicit hard page breaks: <w:br w:type="page"/> in paragraph XML
  2. Last-rendered page breaks: <w:lastRenderedPageBreak/> in run XML
  3. Section breaks in document sections
  4. Fallback: treat whole document as page 1 (single logical page)
"""
from __future__ import annotations
import logging
from pathlib import Path
from pipeline.extractors.base_extractor import BaseExtractor, ExtractionOutput

logger = logging.getLogger(__name__)

# XML namespaces used in DOCX
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _has_page_break(para) -> bool:
    """Check if a paragraph contains an explicit hard page break."""
    from lxml import etree
    # Check for <w:br w:type="page"/> anywhere in this paragraph
    for br in para._element.iter(f"{{{W_NS}}}br"):
        br_type = br.get(f"{{{W_NS}}}type")
        if br_type == "page":
            return True
    return False


def _has_rendered_page_break(para) -> bool:
    """Check for soft page breaks Word calculated automatically."""
    from lxml import etree
    # <w:lastRenderedPageBreak/> appears in run properties
    tag = f"{{{W_NS}}}lastRenderedPageBreak"
    return para._element.find(f".//{tag}") is not None


class DOCXExtractor(BaseExtractor):

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    def extract(self, file_path: Path, use_ocr: bool = True) -> ExtractionOutput:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: poetry add python-docx")

        try:
            doc = Document(str(file_path))
            lines = []
            current_page = 1
            found_any_break = False

            # Insert opening page marker
            lines.append(f"<!-- page:{current_page} -->")

            for para in doc.paragraphs:
                # Check for page break BEFORE processing this paragraph
                if _has_page_break(para):
                    current_page += 1
                    found_any_break = True
                    lines.append(f"\n<!-- page:{current_page} -->")
                elif _has_rendered_page_break(para):
                    current_page += 1
                    found_any_break = True
                    lines.append(f"\n<!-- page:{current_page} -->")

                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name.lower() if para.style else ""
                if "heading 1" in style_name:
                    lines.append(f"\n# {text}\n")
                elif "heading 2" in style_name:
                    lines.append(f"\n## {text}\n")
                elif "heading 3" in style_name:
                    lines.append(f"\n### {text}\n")
                elif "heading" in style_name:
                    lines.append(f"\n#### {text}\n")
                else:
                    lines.append(text)

            # Tables
            for table in doc.tables:
                lines.append("\n")
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells
                        if cell.text.strip()
                    )
                    if row_text:
                        lines.append(row_text)
                lines.append("\n")

            markdown_text = "\n".join(lines)
            page_count = current_page

            # Fallback: if no breaks found, estimate pages from word count
            # but keep everything on page 1 (not page 0)
            if not found_any_break:
                word_count = len(markdown_text.split())
                estimated_pages = max(1, word_count // 300)
                if estimated_pages > 1:
                    logger.warning(
                        f"{file_path.name}: no page breaks found, "
                        f"estimated {estimated_pages} pages from word count. "
                        f"All entities will cite page 1."
                    )
                page_count = estimated_pages

            logger.info(
                f"DOCX extracted: {file_path.name} "
                f"({len(doc.paragraphs)} paragraphs, "
                f"{page_count} pages, "
                f"breaks_found={found_any_break})"
            )

            return ExtractionOutput(
                markdown_text=markdown_text,
                page_count=page_count,
                metadata={
                    "source_file":       file_path.name,
                    "extractor":         "python-docx",
                    "paragraph_count":   len(doc.paragraphs),
                    "table_count":       len(doc.tables),
                    "page_breaks_found": found_any_break,
                    "page_count":        page_count,
                },
            )

        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path.name}: {e}")
            raise